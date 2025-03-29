from docx2pdf import convert
from docx import Document
import openpyxl
import os

def generate_certificates(template_path, excel_path, output_folder):
    try:
        # Verify files exist
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")
        if not os.path.exists(excel_path):
            raise FileNotFoundError(f"Excel file not found: {excel_path}")

        # Load Excel data
        wb = openpyxl.load_workbook(excel_path)
        ws = wb.active
        
        # Create output folder
        os.makedirs(output_folder, exist_ok=True)
        
        # Process each student
        for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            name = row[0]  # Assuming name is in first column
            
            # Skip empty names
            if not name or not isinstance(name, str):
                print(f"Skipping row {row_idx}: Invalid name - {name}")
                continue
                
            try:
                # Load Word template fresh for each student
                doc = Document(template_path)
                
                # Replace placeholders in all paragraphs
                for paragraph in doc.paragraphs:
                    if '<<NAME>>' in paragraph.text:
                        paragraph.text = paragraph.text.replace('<<NAME>>', name.strip())
                
                # Replace in tables if needed
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if '<<NAME>>' in cell.text:
                                cell.text = cell.text.replace('<<NAME>>', name.strip())
                
                # Save temporary Word file
                safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '_')).strip().replace(' ', '_')
                temp_docx = f"temp_{safe_name}.docx"
                doc.save(temp_docx)
                
                # Convert to PDF
                output_pdf = os.path.join(output_folder, f"Certificate_{safe_name}.pdf")
                convert(temp_docx, output_pdf)
                
                # Clean up temporary file
                os.remove(temp_docx)
                
                print(f"Successfully generated: {output_pdf}")
                
            except Exception as e:
                print(f"Failed to generate certificate for {name} (Row {row_idx}): {str(e)}")
                continue
                
    except Exception as e:
        print(f"Fatal error: {str(e)}")

# Usage with error reporting
print("Starting certificate generation...")
generate_certificates(
    template_path="certificate_template.docx",
    excel_path="students.xlsx",
    output_folder="certificates"
)
print("Process completed. Check the 'certificates' folder.")